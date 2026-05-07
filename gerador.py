"""
Geração de documentos DOCX a partir de template com placeholders {{CAMPO}}.

Estratégia de substituição:
  - Para cada parágrafo, consolida todos os runs num único run (preservando a
    formatação do primeiro run) antes de substituir.  Isso garante que placeholders
    fragmentados entre múltiplos runs (comportamento comum do Word) sejam capturados.
  - A tabela de itens é detectada pela presença de {{ORD}} e populada dinamicamente
    clonando a linha-template para cada item, via manipulação direta do XML.
  - Campos de nível de documento (CNPJ, TOTAL_GERAL, etc.) são substituídos em TODOS
    os parágrafos e células (incluindo cabeçalho/rodapé e células fora da tabela de itens).
  - Fidelidade visual preservada: apenas o texto dos runs é alterado; fontes, tamanhos,
    negrito, espaçamentos e margens permanecem intactos.
"""
import copy
import logging
import os
import re

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


# ---------------------------------------------------------------------------
# Substituição em parágrafos
# ---------------------------------------------------------------------------

def _texto_paragrafo(para) -> str:
    return "".join(run.text for run in para.runs)


def _consolidar_runs(para) -> None:
    """
    Junta o texto de todos os runs no primeiro run, esvaziando os demais.
    Preserva a formatação (bold, italic, font) do primeiro run.
    Necessário porque o Word frequentemente fragmenta um placeholder em vários runs.
    """
    if len(para.runs) <= 1:
        return
    texto_completo = _texto_paragrafo(para)
    para.runs[0].text = texto_completo
    for run in para.runs[1:]:
        run.text = ""


def _substituir_em_paragrafo(para, dados: dict) -> bool:
    texto = _texto_paragrafo(para)
    if "{{" not in texto:
        return False

    campos_encontrados = _PLACEHOLDER_RE.findall(texto)
    if not campos_encontrados:
        return False

    _consolidar_runs(para)

    novo_texto = texto
    for campo in campos_encontrados:
        valor = str(dados.get(campo, f"{{{{{campo}}}}}"))
        novo_texto = novo_texto.replace(f"{{{{{campo}}}}}", valor)

    if para.runs:
        para.runs[0].text = novo_texto
    return True


# ---------------------------------------------------------------------------
# Iteração por todo o documento (exceto linha-template da tabela de itens)
# ---------------------------------------------------------------------------

def _iterar_paragrafos(doc: Document):
    """Gera todos os parágrafos do documento: corpo, cabeçalho e rodapé."""
    yield from doc.paragraphs
    for secao in doc.sections:
        yield from secao.header.paragraphs
        yield from secao.footer.paragraphs


def _substituir_em_celula(celula, campos: dict, ignorar_linha_xml=None) -> None:
    """Substitui em parágrafos diretos e recursivamente em tabelas aninhadas."""
    for para in celula.paragraphs:
        _substituir_em_paragrafo(para, campos)
    for tabela_aninhada in celula.tables:
        for linha in tabela_aninhada.rows:
            if ignorar_linha_xml is not None and linha._tr is ignorar_linha_xml:
                continue
            for cel in linha.cells:
                _substituir_em_celula(cel, campos, ignorar_linha_xml)


def _substituir_tudo(doc: Document, campos: dict, ignorar_linha_xml=None) -> None:
    """
    Substitui placeholders em todos os parágrafos e tabelas (incluindo aninhadas).
    `ignorar_linha_xml`: elemento XML (w:tr) da linha-template de itens — é pulado
    para que seus placeholders sejam preservados para o passo de clonagem.
    """
    for para in _iterar_paragrafos(doc):
        _substituir_em_paragrafo(para, campos)

    for tabela in doc.tables:
        for linha in tabela.rows:
            if ignorar_linha_xml is not None and linha._tr is ignorar_linha_xml:
                continue
            for celula in linha.cells:
                _substituir_em_celula(celula, campos, ignorar_linha_xml)


# ---------------------------------------------------------------------------
# Tabela de itens
# ---------------------------------------------------------------------------

def _encontrar_tabela_itens(doc: Document):
    """
    Retorna (tabela, linha_template) onde linha_template é a linha que contém
    {{ORD}} ou {{ITEM}}.  Retorna (None, None) se não encontrar.
    """
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto = " ".join(p.text for p in celula.paragraphs)
                if "{{ORD}}" in texto or "{{ITEM}}" in texto:
                    return tabela, linha
    return None, None


def _clonar_linha_com_dados(linha_template, dados_item: dict):
    """
    Clona o elemento XML w:tr da linha-template e substitui os placeholders
    diretamente nos elementos w:t, preservando toda a formatação original.
    """
    novo_tr = copy.deepcopy(linha_template._tr)

    for t_elem in novo_tr.findall(".//" + qn("w:t")):
        if t_elem.text and "{{" in t_elem.text:
            texto = t_elem.text
            for campo, valor in dados_item.items():
                texto = texto.replace(f"{{{{{campo}}}}}", str(valor))
            t_elem.text = texto

    return novo_tr


def _popular_tabela_itens(tabela, linha_template, itens: list[dict]) -> None:
    """
    Remove a linha-template e insere uma linha real por item, na mesma posição.
    """
    tbl = tabela._tbl
    tr_template = linha_template._tr

    # Descobre o índice da linha-template entre todos os w:tr
    todas_tr = tbl.findall(qn("w:tr"))
    idx = todas_tr.index(tr_template)

    # Âncora: linha imediatamente anterior (ou None se for a primeira)
    ancora = todas_tr[idx - 1] if idx > 0 else None

    # Remove a linha-template
    tbl.remove(tr_template)

    # Insere linhas reais na posição correta
    for i, item in enumerate(itens):
        novo_tr = _clonar_linha_com_dados(linha_template, item)
        if ancora is not None:
            ancora.addnext(novo_tr)
            ancora = novo_tr          # próxima linha inserida depois desta
        else:
            # Era a primeira linha da tabela; insere antes de todas as demais
            tbl.insert(0, novo_tr)
            ancora = novo_tr


# ---------------------------------------------------------------------------
# Funções principais
# ---------------------------------------------------------------------------

def gerar_para_bytes(template_path: str, campos: dict, itens: list[dict]) -> bytes:
    """Gera o documento DOCX em memória e retorna como bytes (para download no Streamlit)."""
    import io
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    doc = Document(template_path)
    tabela_itens, linha_template = _encontrar_tabela_itens(doc)
    linha_template_xml = linha_template._tr if tabela_itens else None

    _substituir_tudo(doc, campos, ignorar_linha_xml=linha_template_xml)

    if tabela_itens is not None and itens:
        _popular_tabela_itens(tabela_itens, linha_template, itens)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def gerar_documento(
    template_path: str,
    output_path: str,
    campos: dict,
    itens: list[dict],
) -> None:
    """
    Gera um documento DOCX a partir do template, substituindo placeholders.

    Args:
        template_path: caminho para o .docx template com {{PLACEHOLDERS}}.
        output_path:   caminho de destino do documento gerado.
        campos:        dicionário com os campos de nível de documento.
        itens:         lista de dicionários com os campos de cada linha da tabela.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    doc = Document(template_path)

    # Localiza tabela de itens antes de qualquer substituição
    tabela_itens, linha_template = _encontrar_tabela_itens(doc)

    if tabela_itens is None:
        logger.warning(
            "Tabela de itens não encontrada no template (esperado: célula com {{ORD}} ou {{ITEM}})."
        )
        linha_template_xml = None
    else:
        linha_template_xml = linha_template._tr

    # Substitui todos os campos de nível de documento, pulando a linha-template
    _substituir_tudo(doc, campos, ignorar_linha_xml=linha_template_xml)

    # Popula tabela de itens
    if tabela_itens is not None and itens:
        _popular_tabela_itens(tabela_itens, linha_template, itens)
    elif not itens:
        logger.warning("Nenhum item para inserir na tabela.")

    # Cria diretório de saída se não existir
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    logger.info("Documento salvo: %s", output_path)
